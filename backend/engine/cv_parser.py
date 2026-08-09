"""Extract plain text from CV/resume uploads (PDF, DOCX, TXT)."""

from pathlib import Path


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".txt":
        return data.decode("utf-8", errors="replace").strip()
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext in (".docx", ".doc"):
        return _extract_docx(data)
    raise ValueError(f"Unsupported file type: {ext or 'unknown'}. Use PDF, DOCX, or TXT.")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        from io import BytesIO

        reader = PdfReader(BytesIO(data))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n".join(parts).strip()
    except ImportError as e:
        raise ValueError("PDF support requires pypdf. Install with: pip install pypdf") from e


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except ImportError as e:
        raise ValueError("DOCX support requires python-docx. Install with: pip install python-docx") from e
